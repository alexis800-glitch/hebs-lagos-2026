# -*- coding: utf-8 -*-
"""Apply the approved HEBS Lagos 2026 concept note revisions to the pristine baseline
DOCX, preserving the logo, styles, tables and every untouched section.

Reads : HEBS_Lagos_2026_Concept_Note_baseline.docx (this directory, never edited)
Writes: HEBS_Lagos_2026_Concept_Note.docx (repository root)

See README.md in this directory. Run export_pdf.py afterwards to produce the PDF.
"""
import io, sys, copy
from pathlib import Path
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[2]
SRC = HERE / 'HEBS_Lagos_2026_Concept_Note_baseline.docx'
OUT = REPO / 'HEBS_Lagos_2026_Concept_Note.docx'

d = Document(str(SRC))
kids = list(d.element.body.iterchildren())          # stable snapshot; insertions don't shift it

def P(i):   return Paragraph(kids[i], d)
def T(i):   return Table(kids[i], d)

# ---------- run-property templates harvested from the document itself ----------
BODY_N = copy.deepcopy(P(8).runs[0]._r.find(qn('w:rPr')))        # Calibri 10.5pt 2B2B2B
BODY_B = copy.deepcopy(P(27).runs[1]._r.find(qn('w:rPr')))       # bold 1A1A1A
_t159  = T(159)
CELL_N = copy.deepcopy(_t159.rows[1].cells[1].paragraphs[0].runs[0]._r.find(qn('w:rPr')))
CELL_H = copy.deepcopy(_t159.rows[0].cells[0].paragraphs[0].runs[0]._r.find(qn('w:rPr')))

def _mkrun(rpr, text):
    from docx.oxml import OxmlElement
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r

def segs(p_el, segments, rn=None, rb=None):
    """Rebuild a paragraph's runs from (text, bold) segments, keeping its pPr."""
    rn = BODY_N if rn is None else rn
    rb = BODY_B if rb is None else rb
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    for text, bold in segments:
        p_el.append(_mkrun(rb if bold else rn, text))

def single(p_el, text):
    """Keep run 0 (and its formatting), set its text, drop the rest. Headings/labels."""
    runs = p_el.findall(qn('w:r'))
    assert runs, 'no runs to retext'
    for r in runs[1:]:
        p_el.remove(r)
    ts = runs[0].findall(qn('w:t'))
    for extra in ts[1:]:
        runs[0].remove(extra)
    ts[0].set(qn('xml:space'), 'preserve')
    ts[0].text = text

def cell_set(cell, segments, header=False):
    p_el = cell.paragraphs[0]._p
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    rb = copy.deepcopy(CELL_N)
    if rb.find(qn('w:b')) is not None:
        rb.remove(rb.find(qn('w:b')))
    from docx.oxml import OxmlElement
    b = OxmlElement('w:b'); rb.insert(1, b)
    segs(p_el, segments, rn=CELL_H if header else CELL_N, rb=CELL_H if header else rb)

def clone_after(template_el, anchor_el):
    new = copy.deepcopy(template_el)
    anchor_el.addnext(new)
    return new

def clone_before(template_el, anchor_el):
    new = copy.deepcopy(template_el)
    anchor_el.addprevious(new)
    return new

BODY_TPL = copy.deepcopy(kids[8])      # plain body paragraph
LIST_TPL = copy.deepcopy(kids[35])     # ListNumber item
SEC_TPL  = copy.deepcopy(kids[55])     # section heading with gold rule
SUB_TPL  = copy.deepcopy(kids[57])     # category-level subheading (Georgia 12pt)
GOLD_TPL = copy.deepcopy(kids[49])     # gold bold label
META_TPL = copy.deepcopy(kids[69])     # short italic meta line
BLANK_TPL = copy.deepcopy(kids[87])    # empty spacer paragraph
NAME_TPL = copy.deepcopy(kids[61])     # competition name (Calibri bold 11pt)
TBL3_TPL = copy.deepcopy(kids[86])     # 3-column table
TBL2_TPL = copy.deepcopy(kids[159])    # 2-column table


def cell_sub(cell, old, new):
    """Replace text inside a cell run, leaving every run property untouched."""
    for p in cell.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new, 1)
                return
    raise AssertionError(f'{old!r} not found in cell {cell.text!r}')


def drop(first, last):
    """Delete baseline body elements first..last inclusive."""
    for el in kids[first:last + 1]:
        el.getparent().remove(el)


def build_table(template_el, anchor_el, rows, bold_rows=()):
    """Clone a table template, resize to len(rows), and fill it."""
    new = copy.deepcopy(template_el)
    anchor_el.addnext(new)
    tbl = Table(new, d)
    while len(tbl.rows) > len(rows):
        new.remove(tbl.rows[-1]._tr); tbl = Table(new, d)
    while len(tbl.rows) < len(rows):
        new.append(copy.deepcopy(tbl.rows[-1]._tr)); tbl = Table(new, d)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell_set(tbl.rows[ri].cells[ci], [(val, 0)], header=(ri == 0 or ri in bold_rows))
    return new

# ============================ EVENT OVERVIEW ============================
segs(kids[6], [
    ("The Hair Education Beauty Summit (HEBS) is an international platform for beauty, culture, and creative enterprise. HEBS launched in ", 0),
    ("New Jersey, USA in 2024", 1),
    (" and returned in ", 0),
    ("2025 with a sold-out edition", 1),
    (" that established its reputation internationally. Following those earlier United States editions, HEBS staged another sold-out edition from ", 0),
    ("2–4 May 2026", 1),
    (".", 0)])

segs(kids[7], [
    ("In October 2026 the summit comes home. ", 0),
    ("HEBS Lagos 2026", 1),
    (" will mark the event’s fourth edition and its Lagos debut: three days of education, competition, exhibition, and cultural programming, hosted at two venues on the Lekki-Epe corridor of Lagos.", 0)])

_p = clone_after(BODY_TPL, kids[7])
segs(_p, [
    ("The Lagos edition is broader than the editions before it. Alongside hair, barbering, braiding, makeup, nails, and fashion, the 2026 programme adds ", 0),
    ("cultural food and live music competition", 1),
    (", and is built in partnership with the tourism and hospitality identity of the Epe region that hosts it.", 0)])

segs(kids[8], [("The summit brings together the full value chain of the beauty and creative economy under one roof, including hairstylists, barbers, braiders, makeup artists, nail technicians, fashion designers, chefs, caterers, seafood and cultural food vendors, vocalists and music creatives, educators, salon owners, manufacturers, distributors, investors, tourism and hospitality stakeholders, and brand leaders from Nigeria, across Africa, and the diaspora.", 0)])

segs(kids[9], [
    ("It is where ", 0),
    ("education meets opportunity, creativity meets commerce, culture meets industry, and Nigerian talent connects with the global marketplace.", 1)])

# ============================ AT A GLANCE ============================
g = T(11)
assert g.rows[3].cells[0].text.strip() == 'Total prize pool', g.rows[3].cells[0].text
assert g.rows[4].cells[0].text.strip() == 'Competition tracks', g.rows[4].cells[0].text
cell_set(g.rows[3].cells[0], [("Total prize pool", 0)])
cell_set(g.rows[3].cells[1], [("$87,500 USD / ₦122,500,000", 1)])
cell_set(g.rows[4].cells[0], [("Competition programme", 0)])
cell_set(g.rows[4].cells[1], [("Thirteen competitions across seven categories: barbering, braiding, hair installation and styling, beauty, fashion, culinary and culture, and music and live entertainment", 0)])
cell_set(g.rows[5].cells[1], [("Education, exhibition, competition, panels, networking, cultural food, live music, and hospitality experiences", 0)])

# ============================ DATE AND VENUE ============================
segs(kids[22], [("Exhibition floor, education sessions, panel discussions, and seven competitions: Battle of the Fades, Fast & Flawless Braiding, Taste of Culture, Freestyle Braid Art, Gilded Heritage Nail Art, Freestyle Design, and Mic Drop Live Vocalist.", 0)])
segs(kids[26], [("Exhibition continues, alongside masterclasses and panels, with seven competitions: Taste of Culture, Traditional Braiding, Roots to Royalty Fashion Runway, Bridal Beauty Makeup, Lash Couture, Flawless Frontals, and Fast & Flawless Barber.", 0)])
segs(kids[27], [
    ("Exhibition days are ", 0),
    ("October 24 and 25, 2026", 1),
    (". The Main Summit venue offers three activated floors, a main stage, masterclass spaces, a cultural food and tasting area, and a dedicated exhibition and vendor floor.", 0)])

# competition schedule tables, exactly as timed in the organiser-supplied document
SCHEDULE = [
    ('Saturday, October 24, 2026', [
        ('Battle of the Fades', '12:30 to 1:00 PM', '30 min'),
        ('Fast & Flawless Braiding Competition', '12:30 to 1:00 PM', '30 min'),
        ('Taste of Culture Food Tasting Competition', '12:00 Noon to 5:00 PM', '5 hours'),
        ('Freestyle Braid Art Competition', '1:15 to 2:30 PM', '75 min'),
        ('Gilded Heritage Nail Art Competition', '2:00 to 3:15 PM', '75 min'),
        ('Freestyle Design Competition', '2:45 to 3:45 PM', '60 min'),
        ('Mic Drop Live Vocalist Competition', '4:00 to 5:15 PM', '75 min'),
    ]),
    ('Sunday, October 25, 2026', [
        ('Taste of Culture Food Tasting Competition', '11:00 AM to 4:00 PM', '5 hours'),
        ('Traditional Braiding Competition', '11:30 AM to 12:30 PM', '60 min'),
        ('Roots to Royalty Fashion Runway Competition', '12:45 to 1:30 PM', '45 min'),
        ('Bridal Beauty Makeup Competition', '1:30 to 3:00 PM', '90 min'),
        ('Lash Couture Competition', '1:45 to 2:30 PM', '45 min'),
        ('Flawless Frontals Competition', '2:45 to 3:45 PM', '60 min'),
        ('Fast & Flawless Barber Competition', '4:00 to 4:15 PM', '15 min'),
    ]),
]
_sched_anchor = clone_after(SUB_TPL, kids[27]); single(_sched_anchor, 'Competition Schedule')
for day, rows in SCHEDULE:
    lab = clone_after(GOLD_TPL, _sched_anchor); single(lab, day)
    tb = build_table(TBL3_TPL, lab, [('Competition', 'Time', 'Duration')] + rows)
    _sched_anchor = clone_after(BLANK_TPL, tb)

# ============================ THEME AND POSITIONING ============================
segs(kids[32], [("HEBS Lagos is built as a platform for Nigeria’s beauty and creative industries, positioning the country as a serious destination for beauty, culture, entrepreneurship, and international collaboration.", 0)])
segs(kids[33], [("The 2026 edition widens that platform. Beauty, fashion, hair, food, music, tourism, and hospitality are not separate economies in Nigeria. They are the same creative economy, powered by the same entrepreneurs, serving the same audience, and carrying the same cultural story. HEBS Lagos 2026 programmes them together: the salon floor, the kitchen, the runway, and the stage in one venue, across one weekend.", 0)])
_p = clone_after(BODY_TPL, kids[33])
segs(_p, [("The 2026 edition is also a homecoming. An industry built and proven in the diaspora returns to the culture that shaped it, bringing international standards, international buyers, and international attention with it, and doing so in a region of Lagos whose own identity is built on tourism, hospitality, fishing, and food.", 0)])

# ============================ WHAT SETS HEBS APART ============================
segs(kids[35], [
    ("Industry fusion.", 1),
    (" A single platform where hairstylists, barbers, braiders, makeup artists, nail technicians, fashion designers, chefs, cultural food vendors, vocalists, and educators work alongside each other and alongside global influencers, rather than in separate events.", 0)])
segs(kids[37], [
    ("A serious competition purse.", 1),
    (" An ", 0),
    ("$87,500 USD (₦122,500,000)", 1),
    (" cumulative prize pool across thirteen competitions, the largest in African beauty competition, staged as a full production with music, lighting, and a live audience.", 0)])
_p = clone_after(LIST_TPL, kids[37])
segs(_p, [
    ("Culture on the programme, not on the poster.", 1),
    (" Local food, seafood heritage, and live music are competition categories in their own right, with their own judges, their own stage time, and their own prizes.", 0)])

# ============================ OBJECTIVES ============================
segs(kids[40], [("HEBS Lagos 2026 is built to deliver against seven objectives.", 0)])
segs(kids[42], [
    ("Connect professionals worldwide.", 1),
    (" Create a hub where stylists, barbers, braiders, makeup artists, nail technicians, chefs, music creatives, educators, and brand leaders can meet and build lasting commercial relationships.", 0)])
segs(kids[44], [
    ("Drive commercial growth.", 1),
    (" Connect exhibitors, manufacturers, distributors, and food and hospitality vendors directly to a concentrated, qualified audience and to new distribution channels.", 0)])
obj_new = _p = clone_after(LIST_TPL, kids[44])
segs(_p, [
    ("Celebrate and commercialise local culture.", 1),
    (" Give the food, seafood, music, and hospitality traditions of the host region a professional competitive platform, and turn cultural pride into visibility, bookings, and business for local vendors and creatives.", 0)])
segs(kids[45], [
    ("Position Lagos internationally.", 1),
    (" Establish the city, and the Lekki-Epe tourism corridor within it, as an emerging beauty, culture, and destination-event location that attracts delegates, media, and investment.", 0)])
segs(kids[46], [
    ("Support the next generation.", 1),
    (" Help keep the future of beauty, grooming, fashion, food, and creative enterprise in Nigeria inclusive, innovative, and grounded in technical skill.", 0)])

# ============================ TARGET AUDIENCE ============================
segs(kids[48], [
    ("Expected attendance is 3,000+", 1),
    (" beauty, culture, and creative professionals, educators, creators, and consumers from Nigeria, West Africa, and the wider HEBS community, including international competitors and delegates from the United States, the United Kingdom, and across West Africa.", 0)])

anchor = kids[50]
for label, body in [
    ("Food and Hospitality Attendees",
     "Food vendors, chefs, caterers, seafood vendors, cultural food entrepreneurs, restaurateurs, catering businesses, and hospitality operators."),
    ("Music and Performance Attendees",
     "Vocalists, music creatives, performers, producers, artist managers, and entertainment professionals."),
]:
    lab = clone_after(GOLD_TPL, anchor); single(lab, label)
    bod = clone_after(BODY_TPL, lab);    segs(bod, [(body, 0)])
    anchor = bod

segs(kids[52], [("Global beauty brands, investors, manufacturers, distributors, brand founders, creative entrepreneurs, tourism and hospitality stakeholders, and industry media.", 0)])
single(kids[53], "Community and Cultural Audience")
segs(kids[54], [("The Epe community and its cultural representatives, beauty and food consumers, cultural attendees, and the wider Lagos creative community, drawn by the competition production, runway presentations, cultural food programming, and pre-party programming.", 0)])

# ==================== NEW SECTION: EPE CULTURAL AND TOURISM ACTIVATION ====================
head = clone_before(SEC_TPL, kids[55]); single(head, "Epe Cultural and Tourism Activation")
epe = [
    [("HEBS Lagos 2026 is hosted on the Lekki-Epe corridor, in a part of Lagos State known for its tourism and hospitality economy, its fishing communities, and one of the strongest seafood and cultural food traditions in the country. The 2026 edition is built to reflect that identity rather than sit apart from it.", 0)],
    [("Two new competition categories carry the activation. The ", 0),
     ("Taste of Culture Food Tasting Competition", 1),
     (" puts Epe’s seafood and cultural cooking on a judged stage in front of thousands of attendees, and the ", 0),
     ("Mic Drop Vocalist Competition", 1),
     (" gives local and visiting vocal talent a professional main-stage platform. Both categories are held on the exhibition days, in the same venue, in front of the same audience, judges, sponsors, and media as the beauty championships.", 0)],
    [("The intent is straightforward. Cooks, caterers, seafood vendors, cultural food entrepreneurs, vocalists, and music creatives receive the same standard of platform, judging, production, and commercial exposure that HEBS already gives to stylists, barbers, and braiders.", 0)],
    [("A proposed Epe community sponsorship will support ", 0),
     ("up to 200 contestants", 1),
     (" by covering their entry participation, creating access for local talent and strengthening community participation. Contestant allocation across categories will be confirmed with the sponsor ahead of registration.", 0)],
    [("Beyond competition, the activation is designed to deliver measurable local value: visibility for Epe as a destination for food, hospitality, and cultural tourism; trade for local food vendors, caterers, transport operators, and hospitality businesses across the event weekend; a route to market for cultural food entrepreneurs seeking catering contracts, retail listings, and media coverage; and a recurring annual platform that returns to the region each year rather than a single visit.", 0)],
]
anchor = head
for s in epe:
    anchor = clone_after(BODY_TPL, anchor); segs(anchor, s)

# ============================ COMPETITION CATEGORIES ============================
# Rebuilt from "Hebs Lagos Competitions. reviewed.docx" (organiser-supplied, 2026-08-03):
# thirteen competitions across seven categories. Descriptions are the supplied
# document's own wording. Times/durations come from its schedule tables. Prize
# ceilings come from the linked competition flyers on Google Drive.
#
# Entry fees no longer come from the flyers. The management price revision of
# 2026-08-03 sets one flat contestant fee of ₦50,000 for every competition, published
# in naira only, with no USD equivalent. This mirrors ENTRY_FEE_NGN in lib/competitions.ts.
ENTRY_FEE = '₦50,000'

CATEGORIES = [
    ('Barbering', [
        ('Battle of the Fades Competition',
         'Saturday, October 24, 2026, 12:30 to 1:00 PM, 30 minutes.',
         ENTRY_FEE, 'up to $5,000 USD (₦7,000,000)',
         'Battle of the Fades focuses on precision, control, and clean barbering execution. Barbers complete a full fade with a properly cut and styled top, demonstrating seamless blending, sharp detailing, balanced structure, and a polished finish. Each haircut should complement the model’s head shape while showcasing strong technical skill and professional presentation.'),
        ('Freestyle Design Competition',
         'Saturday, October 24, 2026, 2:45 to 3:45 PM, 60 minutes.',
         ENTRY_FEE, 'up to $7,500 USD (₦10,500,000)',
         'Freestyle Design explores the intersection of barbering, creativity, and visual art. Barbers create an original haircut design using clean lines, patterns, shapes, and precision detailing. The finished look should be imaginative, balanced, technically clean, and visually impactful from every angle.'),
        ('Fast & Flawless Barber Competition',
         'Sunday, October 25, 2026, 4:00 to 4:15 PM, 15 minutes.',
         # Flat figure, not a ceiling — management decision. No "up to" prefix.
         ENTRY_FEE, '$5,000 USD (₦7,000,000)',
         'The Fast & Flawless Barber Competition is a timed barbering challenge designed to showcase speed, precision, technical control, and professional execution under intense time pressure. Competitors must complete a professional-quality full haircut on one live model within 15 minutes. The finished look must include a clean, consistent full fade extending around the head and over the occipital bone, a sharp and balanced lineup, a properly cut and styled top, and a polished, competition-ready finish.'),
    ]),
    ('Braiding', [
        ('Fast & Flawless Braiding Competition',
         'Saturday, October 24, 2026, 12:30 to 1:00 PM, 30 minutes.',
         ENTRY_FEE, 'up to $5,000 USD (₦7,000,000)',
         'Fast & Flawless Braiding focuses on speed, precision, and natural-hair technique. Braiders complete eight full-head cornrows within 30 minutes using only the model’s natural hair. The finished design should demonstrate clean parting, consistent braid size, even spacing, controlled tension, and polished execution.'),
        ('Freestyle Braid Art Competition',
         'Saturday, October 24, 2026, 1:15 to 2:30 PM, 75 minutes.',
         ENTRY_FEE, 'up to $10,000 USD (₦14,000,000)',
         'Freestyle Braid Art explores creativity, innovation, and technical mastery through braiding. Braiders transform natural hair into intricate works of braid art using original patterns, dimensional structure, creative direction, and artistic detail. The finished design should be cohesive, technically clean, visually impactful, and recognizable as a braid-driven work of art.'),
        ('Traditional Braiding Competition',
         'Sunday, October 25, 2026, 11:30 AM to 12:30 PM, 60 minutes.',
         ENTRY_FEE, 'up to $7,500 USD (₦10,500,000)',
         'Traditional Braiding celebrates the artistry, cultural significance, and technical mastery of African braiding traditions. Braiders create a polished natural-hair design inspired by traditional African patterns, techniques, and cultural influence. The finished style should demonstrate precise sectioning, consistent technique, thoughtful design, and respect for traditional braiding artistry.'),
    ]),
    ('Hair Installation & Styling', [
        ('Flawless Frontals Competition',
         'Sunday, October 25, 2026, 2:45 to 3:45 PM, 60 minutes.',
         ENTRY_FEE, 'up to $5,000 USD (₦7,000,000)',
         'Flawless Frontals focuses on creating a seamless, natural-looking frontal installation. Stylists customize, place, secure, blend, and style the frontal while complementing the model’s features and desired look. The finished installation should demonstrate a refined hairline, clean application, balanced styling, and a polished finish from every angle.'),
    ]),
    ('Beauty', [
        ('Bridal Beauty Makeup Competition',
         'Sunday, October 25, 2026, 1:30 to 3:00 PM, 90 minutes.',
         ENTRY_FEE, 'up to $7,500 USD (₦10,500,000)',
         'Bridal Beauty focuses on creating a radiant, elegant, wedding-ready look. Artists present their interpretation of the modern bride while complementing the model’s features, skin tone, and personal beauty, anywhere from soft and timeless to bold and glamorous. Judged on complexion work, balanced colour, clean application, attention to detail, and long-lasting wear.'),
        ('Lash Couture Competition',
         'Sunday, October 25, 2026, 1:45 to 2:30 PM, 45 minutes.',
         ENTRY_FEE, 'up to $5,000 USD (₦7,000,000)',
         'Lash Couture focuses on creating a customized, polished lash look that enhances the model’s natural eye shape. Artists demonstrate precise isolation, controlled placement, symmetry, balance, and thoughtful lash mapping. The finished set should be clean, comfortable, technically sound, and professionally finished.'),
        ('Gilded Heritage Nail Art Competition',
         'Saturday, October 24, 2026, 2:00 to 3:15 PM, 75 minutes.',
         ENTRY_FEE, 'up to $7,500 USD (₦10,500,000)',
         'Gilded Heritage explores culture, craftsmanship, and luxury through nail design. Artists create a complete set inspired by heritage, traditional patterns, precious metals, textiles, jewellery, and cultural detail, using rich colour, gold accents, sculpted elements, intricate pattern, embellishment, and creative texture. The finished set should be original, balanced, technically clean, and a modern take on culture and elegance.'),
    ]),
    ('Fashion', [
        ('Roots to Royalty Fashion Runway Competition',
         'Sunday, October 25, 2026, 12:45 to 1:30 PM, 45 minutes.',
         ENTRY_FEE, '$7,500 USD (₦10,500,000)',
         'Roots to Royalty celebrates the journey from cultural heritage to modern-day royalty. Designers create an original, runway-ready look inspired by ancestry, legacy, leadership, and the spirit of Lagos, using regal silhouettes, bold colour, rich fabric, cultural influence, symbolic detail, and statement accessories. Each look should honour its roots while presenting a fresh, elevated take on royalty, with cultural inspiration represented thoughtfully and respectfully.'),
    ]),
    ('Culinary & Culture', [
        ('Taste of Culture Food Tasting Competition',
         'Saturday, October 24, 2026, 12:00 Noon to 5:00 PM and Sunday, October 25, 2026, 11:00 AM to 4:00 PM, five hours each day.',
         ENTRY_FEE, 'up to $10,000 USD (₦14,000,000)',
         'Taste of Culture celebrates heritage, community, and creativity through food. Competitors present a dish inspired by their culture, family traditions, regional influence, or personal story. Each entry should demonstrate balanced flavour, thoughtful presentation, quality preparation, and a meaningful connection to the culture it represents.'),
    ]),
    ('Music & Live Entertainment', [
        ('Mic Drop Live Vocalist Competition',
         'Saturday, October 24, 2026, 4:00 to 5:15 PM, 75 minutes.',
         ENTRY_FEE, 'up to $5,000 USD (₦7,000,000)',
         'Mic Drop celebrates live vocal talent, individuality, and powerful stage performance. Vocalists present a song that showcases their voice, musical interpretation, emotional connection, and personal artistry. Each performance should demonstrate vocal control, confidence, originality, stage presence, and the ability to connect with a live audience.'),
    ]),
]
ALL_COMPS = [c for _, comps in CATEGORIES for c in comps]
assert len(ALL_COMPS) == 13, f'expected 13 competitions, got {len(ALL_COMPS)}'
assert len(CATEGORIES[0][1]) == 3, 'expected 3 Barbering competitions'
assert all(fee == ENTRY_FEE for _n, _w, fee, _p, _d in ALL_COMPS), 'every entry fee must be the flat naira fee'

segs(kids[56], [
    ("HEBS Lagos 2026 stages ", 0),
    ("thirteen competitions across seven categories", 1),
    (": barbering, braiding, hair installation and styling, beauty, fashion, culinary and culture, and music and live entertainment. The combined prize pool is ", 0),
    ("$87,500 USD (₦122,500,000)", 1),
    (". The contestant entry fee is ", 0),
    ("₦50,000 per competition", 1),
    (", the same for every competition.", 0)])

# replace the entire former Signature / Barber / Braiding block
drop(57, 81)

anchor = kids[56]
for cat_name, comps in CATEGORIES:
    anchor = clone_after(SUB_TPL, anchor); single(anchor, cat_name)
    for name, when, fee, prize, desc in comps:
        anchor = clone_after(NAME_TPL, anchor); single(anchor, name)
        anchor = clone_after(META_TPL, anchor); single(anchor, when)
        anchor = clone_after(BODY_TPL, anchor)
        segs(anchor, [("Entry fee: ", 0), (fee, 1), (". Prize: ", 0), (prize, 1), (" in cash prizes.", 0)])
        anchor = clone_after(BODY_TPL, anchor); segs(anchor, [(desc, 0)])

# ============================ PRIZE POOL ============================
PRIZE_ROWS = [
    ('Battle of the Fades Competition', '$5,000', '₦7,000,000'),
    ('Freestyle Design Competition', '$7,500', '₦10,500,000'),
    ('Fast & Flawless Barber Competition', '$5,000', '₦7,000,000'),
    ('Fast & Flawless Braiding Competition', '$5,000', '₦7,000,000'),
    ('Freestyle Braid Art Competition', '$10,000', '₦14,000,000'),
    ('Traditional Braiding Competition', '$7,500', '₦10,500,000'),
    ('Flawless Frontals Competition', '$5,000', '₦7,000,000'),
    ('Bridal Beauty Makeup Competition', '$7,500', '₦10,500,000'),
    ('Lash Couture Competition', '$5,000', '₦7,000,000'),
    ('Gilded Heritage Nail Art Competition', '$7,500', '₦10,500,000'),
    ('Roots to Royalty Fashion Runway Competition', '$7,500', '₦10,500,000'),
    ('Taste of Culture Food Tasting Competition', '$10,000', '₦14,000,000'),
    ('Mic Drop Live Vocalist Competition', '$5,000', '₦7,000,000'),
]
_usd = sum(int(r[1].replace('$', '').replace(',', '')) for r in PRIZE_ROWS)
_ngn = sum(int(r[2].replace('₦', '').replace(',', '')) for r in PRIZE_ROWS)
assert _usd == 87500, f'prize rows total ${_usd:,}, expected $87,500'
assert _ngn == 122500000, f'prize rows total ₦{_ngn:,}, expected ₦122,500,000'
assert len(PRIZE_ROWS) == 13

# headline banner
banner = T(84)
cell_sub(banner.rows[0].cells[0], '$80,000 USD / ₦112,000,000', '$87,500 USD / ₦122,500,000')

# replace the three-track table with the thirteen-competition table
old_pt = kids[86]
new_pt = build_table(TBL3_TPL, old_pt,
                     [('Competition', 'USD', 'NGN')] + PRIZE_ROWS +
                     [('Total prize pool', '$87,500', '₦122,500,000')],
                     bold_rows={14})
old_pt.getparent().remove(old_pt)

# ============================ EXPECTED IMPACT ============================
segs(kids[140], [
    ("HEBS Lagos is not only a competition. It is an ", 0),
    ("economic development platform", 1),
    (" intended to help Nigerian and African talent build recognised brands through education, investment, innovation, tourism, food, manufacturing, and international partnership.", 0)])
segs(kids[142], [("A platform that supports entrepreneurs, creates jobs, and drives business growth across Nigeria’s beauty, food, and creative industries by connecting practitioners to training, capital relationships, and the market access they need to scale.", 0)])
segs(kids[144], [("Attracting international visitors, delegates, competitors, and foreign investment to Lagos, positioning the city and the Lekki-Epe corridor as an emerging beauty, culture, and destination-event location, with direct spillover into hospitality, food service, transport, and local commerce in the host community.", 0)])
segs(kids[146], [("Strengthening international trade by connecting Nigerian manufacturing, distribution, product innovation, and cultural food enterprise to global markets, buyers, and distribution partners.", 0)])
segs(kids[148], [("Raising the technical standard of the sector through strong education and international benchmarking, and building a repeatable annual platform that grows in value for participants, brands, and the host region each year.", 0)])
segs(kids[149], [
    ("HEBS Lagos 2026 is designed to become ", 0),
    ("Nigeria’s premier annual beauty, fashion, culture, and creative industries summit", 1),
    (", and an anchor event on the African and international industry calendar.", 0)])

# ============================ REGISTRATION ============================
rt = T(153)
assert 'Registration Opens' in rt.rows[1].cells[0].text, rt.rows[1].cells[0].text
cell_set(rt.rows[1].cells[2], [
    ("Register at hebseventportal.com/register and pay the entry fee for your chosen competition. Open worldwide to barbers, braiders, frontal and installation stylists, makeup artists, lash artists, nail artists, fashion designers, chefs and cultural food entrepreneurs, vocalists, and beauty, culture, and fashion visionaries.", 0)])

# The former Signature grouping no longer exists, so the stages that referenced it are
# generalised. Wording kept minimal: no new requirement is introduced.
cell_set(rt.rows[2].cells[2], [
    ("Where a video submission is required for a competition, competitors submit a 3-minute MP4 or MOV video to casting@hebslagos.com showing their name, competition, work-in-progress clips, and a before and after transformation. All registrations and submissions must be complete by the deadline.", 0)])
cell_set(rt.rows[3].cells[2], [
    ("All entrants are told whether they have been selected to compete.", 0)])
cell_set(rt.rows[4].cells[0], [("Competition Days", 0)], header=True)
cell_set(rt.rows[4].cells[1], [("October 24 and 25, 2026", 0)], header=True)
cell_set(rt.rows[4].cells[2], [
    ("Competitors present at their scheduled time, as published in the competition schedule.", 0)])

# One flat contestant entry fee across the revised thirteen competitions, set by the
# management price revision of 2026-08-03 and published in naira only. The earlier
# flyer-derived USD tiers ($50 / $75 / $100) are withdrawn and must not reappear.
FEE_ROWS = [(name, fee) for _, comps in CATEGORIES for name, _w, fee, _p, _d in comps]
assert len(FEE_ROWS) == 13
assert all(fee == ENTRY_FEE for _n, fee in FEE_ROWS)

# Replace the three per-track fee tables, and the now-obsolete team-entry note, with
# one table covering all thirteen competitions. No competition in the revised programme
# is contested by a team, so the team-entry paragraph no longer applies.
segs(kids[156], [
    ("Every competition carries the same contestant entry fee of ", 0),
    ("₦50,000", 1),
    (". Competitors pay only for the competitions they enter, and may enter more than one.", 0)])
drop(157, 168)
_ft = build_table(TBL2_TPL, kids[156], [('Competition', 'Entry Fee')] + FEE_ROWS)
clone_after(BLANK_TPL, _ft)
cell_sub_p = kids[169]
segs(cell_sub_p, [
    ("All entry fees are ", 0), ("non-refundable", 1),
    (" and are quoted and paid in naira on the official portal. Competitors arrange and pay for their own travel and lodging to Lagos. HEBS provides the stage production and, where applicable, the official competition equipment and styling products supplied by our partner brands.", 0)])

# ============================ OBJECTIVES LIST RESTART ============================
# The List Number style shares one numId, so Objectives continued the previous list
# (originally 5-10 under "six objectives"). Give Objectives its own restarting numId.
from docx.oxml import OxmlElement

numbering = d.part.numbering_part.element
NEW_NUMID = '10'
_n = OxmlElement('w:num')
_n.set(qn('w:numId'), NEW_NUMID)
_a = OxmlElement('w:abstractNumId'); _a.set(qn('w:val'), '7'); _n.append(_a)   # same look as List Number
_ov = OxmlElement('w:lvlOverride'); _ov.set(qn('w:ilvl'), '0')
_so = OxmlElement('w:startOverride'); _so.set(qn('w:val'), '1'); _ov.append(_so)
_n.append(_ov)
numbering.append(_n)

def set_numid(p_el, numid):
    pPr = p_el.find(qn('w:pPr'))
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = OxmlElement('w:numPr')
        pStyle = pPr.find(qn('w:pStyle'))
        (pStyle.addnext(numPr) if pStyle is not None else pPr.insert(0, numPr))
    for tag in ('w:ilvl', 'w:numId'):
        e = numPr.find(qn(tag))
        if e is not None:
            numPr.remove(e)
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    nid = OxmlElement('w:numId'); nid.set(qn('w:val'), numid)
    numPr.append(ilvl); numPr.append(nid)

for _el in [kids[41], kids[42], kids[43], kids[44], obj_new, kids[45], kids[46]]:
    set_numid(_el, NEW_NUMID)

d.save(str(OUT))
print('saved', OUT)
